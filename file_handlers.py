"""
Enhanced file handlers for multiple file formats
"""
import os
import json
import csv
import pandas as pd
import numpy as np
from docx import Document
from io import BytesIO, StringIO
import PyPDF2
import pdfplumber
from bs4 import BeautifulSoup
import chardet
import markdown
from typing import Optional, Tuple, Dict, Any
import logging
import re

logger = logging.getLogger(__name__)

class FileHandler:
    """Universal file handler for multiple formats"""
    
    SUPPORTED_FORMATS = {
        'text/plain': ['.txt', '.text', '.log', '.md', '.markdown'],
        'application/pdf': ['.pdf'],
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['.docx'],
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': ['.xlsx'],
        'application/vnd.ms-excel': ['.xls'],
        'text/csv': ['.csv'],
        'application/json': ['.json'],
        'text/html': ['.html', '.htm'],
        'application/xml': ['.xml'],
        'application/rtf': ['.rtf'],
        'text/tab-separated-values': ['.tsv']
    }
    
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    
    @classmethod
    def get_supported_extensions(cls) -> list:
        """Get all supported file extensions"""
        extensions = []
        for ext_list in cls.SUPPORTED_FORMATS.values():
            extensions.extend(ext_list)
        return extensions
    
    @classmethod
    def validate_file(cls, file, max_size: int = None) -> Tuple[bool, str]:
        """Validate uploaded file"""
        if not file:
            return False, "No file provided"
        
        # Check file extension
        file_ext = os.path.splitext(file.name)[1].lower()
        if file_ext not in cls.get_supported_extensions():
            return False, f"Unsupported file format: {file_ext}"
        
        # Check file size
        max_size = max_size or cls.MAX_FILE_SIZE
        file.seek(0, 2)
        size = file.tell()
        file.seek(0)
        
        if size > max_size:
            return False, f"File size ({size/1024/1024:.1f}MB) exceeds maximum ({max_size/1024/1024:.1f}MB)"
        
        return True, "File valid"
    
    @classmethod
    def process_file(cls, file) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
        """Process any supported file format"""
        try:
            file_ext = os.path.splitext(file.name)[1].lower()
            file_content = file.read()
            
            # Reset file pointer
            file.seek(0)
            
            # Route to appropriate handler
            if file_ext in ['.txt', '.text', '.log']:
                return cls._process_text_file(file_content)
            elif file_ext in ['.md', '.markdown']:
                return cls._process_markdown_file(file_content)
            elif file_ext == '.pdf':
                return cls._process_pdf_file(file)
            elif file_ext == '.docx':
                return cls._process_docx_file(file_content)
            elif file_ext in ['.xlsx', '.xls']:
                return cls._process_excel_file(file_content, file_ext)
            elif file_ext == '.csv':
                return cls._process_csv_file(file_content)
            elif file_ext == '.tsv':
                return cls._process_tsv_file(file_content)
            elif file_ext == '.json':
                return cls._process_json_file(file_content)
            elif file_ext in ['.html', '.htm']:
                return cls._process_html_file(file_content)
            elif file_ext == '.xml':
                return cls._process_xml_file(file_content)
            elif file_ext == '.rtf':
                return cls._process_rtf_file(file_content)
            else:
                return None, {"error": f"Unsupported file format: {file_ext}"}
                
        except Exception as e:
            logger.error(f"Error processing file {file.name}: {e}")
            return None, {"error": str(e)}
    
    @staticmethod
    def _process_text_file(content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Process plain text files"""
        # Detect encoding
        detected = chardet.detect(content)
        encoding = detected.get('encoding', 'utf-8')
        
        try:
            text = content.decode(encoding, errors='ignore')
        except:
            text = content.decode('utf-8', errors='ignore')
        
        metadata = {
            "type": "text",
            "encoding": encoding,
            "lines": len(text.splitlines()),
            "characters": len(text)
        }
        
        return text, metadata
    
    @staticmethod
    def _process_markdown_file(content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Process Markdown files"""
        text = content.decode('utf-8', errors='ignore')
        
        # Convert markdown to plain text (remove formatting)
        html = markdown.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')
        plain_text = soup.get_text()
        
        # Extract metadata
        headers = re.findall(r'^#{1,6}\s+(.+)$', text, re.MULTILINE)
        links = re.findall(r'\[([^\]]+)\]\(([^\)]+)\)', text)
        
        metadata = {
            "type": "markdown",
            "headers": headers,
            "links": len(links),
            "has_code_blocks": '```' in text,
            "original_length": len(text),
            "plain_text_length": len(plain_text)
        }
        
        return plain_text, metadata
    
    @staticmethod
    def _process_pdf_file(file) -> Tuple[str, Dict[str, Any]]:
        """Process PDF files with enhanced extraction"""
        text = ""
        metadata = {"type": "pdf", "pages": 0, "has_images": False}
        
        try:
            # Try pdfplumber first (better for tables)
            with pdfplumber.open(file) as pdf:
                metadata["pages"] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num} ---\n{page_text}"
                    
                    # Check for tables
                    tables = page.extract_tables()
                    if tables:
                        metadata["has_tables"] = True
                        for table in tables:
                            # Convert table to text
                            table_text = "\n".join(["\t".join(str(cell) if cell else "" for cell in row) for row in table])
                            text += f"\n[Table on page {page_num}]\n{table_text}\n"
                    
                    # Check for images
                    if hasattr(page, 'images') and page.images:
                        metadata["has_images"] = True
        
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                file.seek(0)
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num} ---\n{page_text}"
            except Exception as e2:
                logger.error(f"Both PDF readers failed: {e2}")
                return None, {"error": f"Failed to read PDF: {str(e2)}"}
        
        metadata["extracted_length"] = len(text)
        return text.strip(), metadata
    
    @staticmethod
    def _process_docx_file(content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Process Word documents with enhanced extraction"""
        doc = Document(BytesIO(content))
        
        text_parts = []
        metadata = {
            "type": "docx",
            "paragraphs": 0,
            "tables": 0,
            "images": 0,
            "headers": [],
            "lists": 0
        }
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
                metadata["paragraphs"] += 1
                
                # Check for headers (based on style)
                if paragraph.style and 'Heading' in paragraph.style.name:
                    metadata["headers"].append(paragraph.text[:100])
                
                # Check for lists
                if paragraph.style and ('List' in paragraph.style.name or 
                                       paragraph.text.strip().startswith(('•', '-', '*', '1.', '2.'))):
                    metadata["lists"] += 1
        
        # Extract tables
        for table in doc.tables:
            metadata["tables"] += 1
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append('\t'.join(row_text))
            text_parts.append('\n[Table]\n' + '\n'.join(table_text))
        
        # Check for images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                metadata["images"] += 1
        
        text = '\n'.join(text_parts)
        metadata["total_length"] = len(text)
        
        return text, metadata
    
    @staticmethod
    def _process_excel_file(content: bytes, ext: str) -> Tuple[str, Dict[str, Any]]:
        """Process Excel files with multiple sheets"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(BytesIO(content))
            
            text_parts = []
            metadata = {
                "type": "excel",
                "sheets": len(excel_file.sheet_names),
                "sheet_names": excel_file.sheet_names,
                "total_rows": 0,
                "total_columns": 0
            }
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # Add sheet header
                text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")
                
                # Convert dataframe to text
                text_parts.append(df.to_string())
                
                # Update metadata
                metadata["total_rows"] += len(df)
                metadata["total_columns"] = max(metadata["total_columns"], len(df.columns))
                
                # Add summary statistics for numeric columns
                numeric_cols = df.select_dtypes(include=[np.number]).columns
                if len(numeric_cols) > 0:
                    text_parts.append("\n[Numeric Summary]")
                    text_parts.append(df[numeric_cols].describe().to_string())
            
            text = '\n'.join(text_parts)
            metadata["total_length"] = len(text)
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing Excel file: {e}")
            return None, {"error": str(e)}
    
    @staticmethod
    def _process_csv_file(content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Process CSV files"""
        try:
            # Detect encoding
            detected = chardet.detect(content[:10000])  # Sample first 10KB
            encoding = detected.get('encoding', 'utf-8')
            
            # Try to read CSV
            text_content = content.decode(encoding, errors='ignore')
            
            # Use StringIO to read CSV
            df = pd.read_csv(StringIO(text_content))
            
            metadata = {
                "type": "csv",
                "encoding": encoding,
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist(),
                "has_header": True
            }
            
            # Convert to text
            text_parts = [df.to_string()]
            
            # Add summary for numeric columns
            numeric_cols = df.select_dtypes(include=[np.number]).columns
            if len(numeric_cols) > 0:
                text_parts.append("\n[Numeric Summary]")
                text_parts.append(df[numeric_cols].describe().to_string())
            
            text = '\n'.join(text_parts)
            metadata["total_length"] = len(text)
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing CSV file: {e}")
            # Try simple text extraction as fallback
            try:
                text = content.decode('utf-8', errors='ignore')
                return text, {"type": "csv", "error": "Processed as plain text", "length": len(text)}
            except:
                return None, {"error": str(e)}
    
    @staticmethod
    def _process_tsv_file(content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Process TSV files"""
        try:
            text_content = content.decode('utf-8', errors='ignore')
            df = pd.read_csv(StringIO(text_content), sep='\t')
            
            metadata = {
                "type": "tsv",
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist()
            }
            
            text = df.to_string()
            metadata["total_length"] = len(text)
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing TSV file: {e}")
            return None, {"error": str(e)}
    
    @staticmethod
    def _process_json_file(content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Process JSON files"""
        try:
            text = content.decode('utf-8', errors='ignore')
            data = json.loads(text)
            
            metadata = {
                "type": "json",
                "structure": type(data).__name__,
                "keys": list(data.keys()) if isinstance(data, dict) else None,
                "length": len(data) if isinstance(data, (list, dict)) else None
            }
            
            # Pretty print JSON for readability
            formatted_text = json.dumps(data, indent=2, ensure_ascii=False)
            
            # Also create a flattened text version for analysis
            def flatten_json(obj, prefix=''):
                items = []
                if isinstance(obj, dict):
                    for k, v in obj.items():
                        new_key = f"{prefix}.{k}" if prefix else k
                        if isinstance(v, (dict, list)):
                            items.extend(flatten_json(v, new_key))
                        else:
                            items.append(f"{new_key}: {v}")
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        new_key = f"{prefix}[{i}]"
                        if isinstance(item, (dict, list)):
                            items.extend(flatten_json(item, new_key))
                        else:
                            items.append(f"{new_key}: {item}")
                else:
                    items.append(f"{prefix}: {obj}")
                return items
            
            flattened = flatten_json(data)
            analysis_text = '\n'.join(flattened)
            
            metadata["flattened_items"] = len(flattened)
            metadata["total_length"] = len(analysis_text)
            
            return analysis_text, metadata
            
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON: {e}")
            # Return raw text if JSON parsing fails
            text = content.decode('utf-8', errors='ignore')
            return text, {"type": "json", "error": "Invalid JSON, processed as text", "length": len(text)}
        except Exception as e:
            logger.error(f"Error processing JSON file: {e}")
            return None, {"error": str(e)}
    
    @staticmethod
    def _process_html_file(content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Process HTML files"""
        try:
            # Detect encoding
            detected = chardet.detect(content[:10000])
            encoding = detected.get('encoding', 'utf-8')
            
            html = content.decode(encoding, errors='ignore')
            soup = BeautifulSoup(html, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            # Extract metadata
            metadata = {
                "type": "html",
                "encoding": encoding,
                "title": soup.title.string if soup.title else None,
                "headers": [h.get_text().strip() for h in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])],
                "links": len(soup.find_all('a')),
                "images": len(soup.find_all('img')),
                "paragraphs": len(soup.find_all('p')),
                "tables": len(soup.find_all('table'))
            }
            
            metadata["total_length"] = len(text)
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing HTML file: {e}")
            return None, {"error": str(e)}
    
    @staticmethod
    def _process_xml_file(content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Process XML files"""
        try:
            xml = content.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(xml, 'xml')
            
            # Extract all text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            text = '\n'.join(line for line in lines if line)
            
            # Extract metadata
            metadata = {
                "type": "xml",
                "root_tag": soup.find().name if soup.find() else None,
                "total_tags": len(soup.find_all()),
                "unique_tags": len(set(tag.name for tag in soup.find_all()))
            }
            
            metadata["total_length"] = len(text)
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing XML file: {e}")
            return None, {"error": str(e)}
    
    @staticmethod
    def _process_rtf_file(content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Process RTF files"""
        try:
            # Basic RTF to text conversion
            rtf = content.decode('utf-8', errors='ignore')
            
            # Remove RTF formatting (simplified)
            # This is a basic implementation - for production use striprtf library
            text = re.sub(r'\\[a-z]+\d*\s?', '', rtf)  # Remove control words
            text = re.sub(r'[{}]', '', text)  # Remove braces
            text = re.sub(r'\\[\\{}]', '', text)  # Remove escaped characters
            text = text.strip()
            
            metadata = {
                "type": "rtf",
                "original_length": len(rtf),
                "text_length": len(text)
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing RTF file: {e}")
            return None, {"error": str(e)}
    
    @classmethod
    def get_file_preview(cls, text: str, max_length: int = 500) -> str:
        """Get a preview of file content"""
        if not text:
            return "No content available"
        
        if len(text) <= max_length:
            return text
        
        # Try to break at sentence boundary
        preview = text[:max_length]
        last_period = preview.rfind('.')
        last_newline = preview.rfind('\n')
        
        break_point = max(last_period, last_newline)
        if break_point > max_length * 0.7:
            preview = text[:break_point + 1]
        
        return preview + "\n\n[... Preview truncated ...]"